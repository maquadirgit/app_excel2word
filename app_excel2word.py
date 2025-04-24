import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import io

def create_tech_doc_from_excel(df):
    doc = Document()
    doc.add_heading("Informatica Mapping - Technical Design Document", 0)

    grouped = df.groupby("Transformation Name")

    for trans_name, group in grouped:
        doc.add_heading(trans_name, level=1)
        trans_type = group["Transformation Type"].iloc[0]
        doc.add_paragraph(f"**Transformation Type:** {trans_type}", style="List Bullet")

        for _, row in group.iterrows():
            field = row["Field"]
            logic = row["Logic"]
            p = doc.add_paragraph()
            p.add_run(f"Field: {field}").bold = True
            p.add_run(f"\nLogic:\n{logic}")
    
    return doc

# Streamlit UI
st.set_page_config(page_title="Excel to Word - Tech Design Generator", layout="centered")
st.title("📄 Excel to Word - Informatica Technical Design Generator")

uploaded_file = st.file_uploader("Upload Excel File", type="xlsx")

if uploaded_file:
    try:
        df = pd.read_excel(uploaded_file)
        required_cols = {"Transformation Type", "Transformation Name", "Field", "Logic"}
        if not required_cols.issubset(df.columns):
            st.error("Uploaded Excel must have columns: Transformation Type, Transformation Name, Field, Logic")
        else:
            st.success("Excel file successfully loaded!")

            doc = create_tech_doc_from_excel(df)

            buffer = io.BytesIO()
            doc.save(buffer)
            st.download_button(
                label="📥 Download Word Document",
                data=buffer.getvalue(),
                file_name="Informatica_Technical_Design.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"Error processing file: {e}")
